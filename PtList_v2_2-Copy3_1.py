#!/usr/bin/env python
# coding: utf-8

# In[ ]:


print('程式載入較久，感謝您耐心等候\n************注意事項************\n1.使用時請確認電腦已連上院內網路\n2.輸入時請仔細確認內容是否正確，\n 再按下Enter，程式中並無防呆機制\n3.程式不正常結束、閃退、印出文件\n 內容有問題等，請參考readme.pdf\n******************by M114@陳郁政')
# ------import---------
import requests
from bs4 import BeautifulSoup
from lxml import html
import datetime
import re
import time
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---函數區---
def GetPtList(Section, Ward, Doc, WardDel, BlockW, BlockBs, BlockBe, sessions) :
    """產生病人清單"""
    # -----------病人資料--------------
    PtList={ 'CHARTNO': [], 'MEDNO' : [], 'SECTIONNO' : [], 'DepName' : [], 'HCASENO' : [], 'INDATETIME' : [], 'NRCODE' : [], 'BEDNO' : [], 'NameGenderAge' : [], 'VSDR' : [], 'VSDRNAME' : [], 'VISITSEQ' : [], 'PtData' : [] }
    # ------------抽取病人清單------------
    for s in Section :
        for w in Ward :
            formdata = {
                'SearchSectionNo' : s,
                'SearchNRCode' : w,
                'SearchVSDR' : Doc,
            }
            list_requ = sessions.post('http://mobilereport.ndmctsgh.edu.tw/eForm/Patient/Result', data = formdata).text
            list_txt = eval(list_requ.replace('null', ' "" '))
            # ---剔除結案即出院病人---
            list_txt = [ txt for txt in list_txt if txt['OUTDATETIME'] == '' and txt['PLCLOSE'] != '1' ]
            # ---剔除指定病房病人---
            list_txt = [ txt for txt in list_txt if txt['NRCODE'] not in WardDel ]
            # ---指定床段---
            list_txt = [ t for t in list_txt if t['NRCODE'] not in BlockW or any([ True for i in range(len(BlockW)) if t['NRCODE'] == BlockW[i] and (t['BEDNO'] >= BlockBs[i] and t['BEDNO'] <= BlockBe[i]) ]) ]
            for txt in list_txt :
                PtList['CHARTNO'].append(txt['CHARTNO'])
                PtList['MEDNO'].append(txt['MEDNO'])
                PtList['SECTIONNO'].append(txt['SECTIONNO'])
                PtList['DepName'].append(txt['DepName'])
                PtList['HCASENO'].append(txt['HCASENO'])
                PtList['INDATETIME'].append(txt['INDATETIME'])
                PtList['NRCODE'].append(txt['NRCODE'])
                PtList['BEDNO'].append(txt['BEDNO'])
                PtList['NameGenderAge'].append(txt['NameGenderAge'])
                PtList['VSDR'].append(txt['VSDR'])
                PtList['VSDRNAME'].append(txt['VSDRNAME'])
                PtList['VISITSEQ'].append(txt['VISITSEQ'])
                PtList['PtData'].append([])
    return PtList
# ------------------------------
def GetTPR(CHARTNO, MEDNO, VISITSEQ, sessions) : 
    """TPRList"""
    print('TPR', end = '->')
    eDate = datetime.date.today()
    sDate = eDate - datetime.timedelta(2)
    eDate = eDate.strftime('%Y-%m-%d')
    sDate = sDate.strftime('%Y-%m-%d')
    load_vitalsign = {
        'ptData[CHARTNO]' : CHARTNO,
        'ptData[MEDNO]' : MEDNO,
        'ptData[VISITSEQ]' : VISITSEQ,
        'sDate' : sDate,
        'eDate' : eDate,
    }

    requ_VitalSign = sessions.post('http://mobilereport.ndmctsgh.edu.tw/eForm/PL/ChangeCare/VitalSignList', data = load_vitalsign)

    data_TPR = { 'date' : [], 'time' : [], 'TPR' : [], }

    Temp = eval('['+re.findall(r"var dataBt = \$\.parseJSON\(\'\[(.*)\]", requ_VitalSign.text)[0].replace('null', ' "" ')+']')
    Temp_list = []
    T_time = []
    for t in Temp : 
        Temp_list.append(t['x']+'T'+str(t['y']))
        T_time.append(t['x'])
    Temp_list = sorted(set(Temp_list))
    T_time = sorted(set(T_time))

    Pulse = eval('['+re.findall(r"var dataPulse = \$\.parseJSON\(\'\[(.*)\]", requ_VitalSign.text)[0].replace('null', ' "" ')+']')
    Pulse_list = []
    P_time = []
    for t in Pulse : 
        Pulse_list.append(t['x']+'T'+str(t['y']))
        P_time.append(t['x'])
    Pulse_list = sorted(set(Pulse_list))
    P_time = sorted(set(P_time))

    Resp = eval('['+re.findall(r"var dataRespiratory = \$\.parseJSON\(\'\[(.*)\]", requ_VitalSign.text)[0].replace('null', ' "" ')+']')
    Resp_list = []
    R_time = []
    for t in Resp : 
        Resp_list.append(t['x']+'T'+str(t['y']))
        R_time.append(t['x'])
    Resp_list = sorted(set(Resp_list))
    R_time = sorted(set(R_time))

    Tc = set(T_time).intersection(set(P_time),set(R_time))
    key_T = [ x for x in range(len(T_time)) if T_time[x] in Tc]
    key_P = [ x for x in range(len(P_time)) if P_time[x] in Tc]
    key_R = [ x for x in range(len(R_time)) if R_time[x] in Tc]

    Temp_list = [ Temp_list[k] for k in key_T ]
    Pulse_list = [ Pulse_list[k] for k in key_P ]
    Resp_list = [ Resp_list[k] for k in key_R ]

    for i in range(len(Temp_list)):
        data_TPR['date'].append(Temp_list[i].split('T')[0])
        data_TPR['time'].append(Temp_list[i].split('T')[1][0:5])
        data_TPR['TPR'].append(Temp_list[i].split('T')[2]+'/'+Pulse_list[i].split('T')[2]+'/'+Resp_list[i].split('T')[2])
    
    TPR = []
    date = list(sorted(set(data_TPR['date'])))
    for i in range(min(3,len(date))) :
        t = [ data_TPR['time'][x]+'-'+data_TPR['TPR'][x] for x in range(len(data_TPR['date'])) if data_TPR['date'][x] == date[i] ]
        if len(t) > 5 :
            t = [ t[int(i*(len(t)-1)/4)] for i in list(range(5)) ]
        TPR.append({
             'date' : date[i],
             'TPR' : t
         })   
    
    return TPR
# ------------------------------
def GetAntiList(CHARTNO, MEDNO, VISITSEQ, sessions) :
    """Anti List"""
    print('anti', end = '->')
    eDate = datetime.date.today()
    sDate = eDate - datetime.timedelta(2)
    eDate = eDate.strftime('%Y-%m-%d')
    sDate = sDate.strftime('%Y-%m-%d')
    load_anti = {
        'ptData[CHARTNO]' : CHARTNO,
        'ptData[MEDNO]' : MEDNO,
        'ptData[VISITSEQ]' : VISITSEQ,
        'sDate' : eDate,
        'eDate' : eDate
    }
    requ_anti = sessions.post('http://mobilereport.ndmctsgh.edu.tw/eForm/PL/ChangeCare/AntibioticList', data = load_anti)

    anti = eval('['+re.findall(r"var paData = \$\.parseJSON\(\'\[(.*)\]", requ_anti.text)[0].replace('null',' "" ')+']')
    anti_list = {
        'name' : [],
        'dose' : [],
        'date' : []
    }
    for a in anti :
        anti_list['name'].append(a['name'].split(' ')[0]+' '+a['name'].split(' ')[1])
        anti_list['dose'].append(str(a['data'][0]['DOSE'])+a['data'][0]['ORDERUNIT']+' '+a['data'][0]['FREQNO'])
        anti_list['date'].append(a['data'][0]['BEGINDATETIME'][3:7]+'-'+a['data'][0]['ENDDATETIME'][3:7])
    return anti_list
# ------------------------------
def GetIOList(CHARTNO, MEDNO, VISITSEQ, sessions) :
    """BP/SpO2/IO List"""
    print('BP/SpO2/IO', end = '->')
    eDate = datetime.date.today()
    sDate = eDate - datetime.timedelta(2)
    eDate = eDate.strftime('%Y-%m-%d')
    sDate = sDate.strftime('%Y-%m-%d')
    load_IO = {
        'ptData[CHARTNO]' : CHARTNO,
        'ptData[MEDNO]' : MEDNO,
        'ptData[VISITSEQ]' : VISITSEQ,
        'sDate' : sDate,
        'eDate' : eDate
    }
    requ_IO = sessions.post('http://mobilereport.ndmctsgh.edu.tw/eForm/PL/ChangeCare/IoList', data = load_IO)
    soup = BeautifulSoup(requ_IO.text, 'html.parser').select('table.table')

    getdate = soup[0].find_all('th')
    IO_list = []
    for d in getdate[1::] :
        IO_list.append(
            {
            'date' : d.text.replace('\u3000',''),
            'BP' : [], 'SpO2' : [],
            'IO' : ['I: -', 'O: -', 'sum: -'], 'IO-I' : [], 'IO-O' : [], 'stool' : [],
            })
    getdata = soup[0].find_all('tr')

    BP = getdata[1].select('td')[1::]
    SpO2 = getdata[2].select('td')[1::]
    stool = getdata[5].select('td')[1::]
    IO = getdata[6].select('td')[1::]
    IO_I = getdata[7].select('td')[1::]
    IO_O = getdata[8].select('td')[1::]

    for i in range(len(IO_list)) :
        for k in range(len(BP[i].select('span.occurDate'))) :
            IO_list[i]['BP'].append(BP[i].select('span.occurDate')[k].text+'-'+BP[i].select('span.nValue')[k].text.replace(' ',''))
        if len(IO_list[i]['BP']) > 5 : 
            IO_list[i]['BP'] = [ IO_list[i]['BP'][int(j*(len(IO_list[i]['BP'])-1)/4)] for j in list(range(5)) ]
        for k in range(len(SpO2[i].select('span.occurDate'))) :
            IO_list[i]['SpO2'].append(SpO2[i].select('span.occurDate')[k].text+'-'+SpO2[i].select('span.nValue')[k].text.replace('SPO2:','').replace(' ',''))
        if len(IO_list[i]['SpO2']) > 5 : 
            IO_list[i]['SpO2'] = [ IO_list[i]['SpO2'][int(j*(len(IO_list[i]['SpO2'])-1)/4)] for j in list(range(5)) ]
        for k in range(len(stool[i].select('span.occurDate'))) :
            IO_list[i]['stool'].append(stool[i].select('span.nValue')[k].text.split(' ')[0])
        if IO[i].select('span.nValue') != [] :
            for j in range(len(IO[i].select('span.nValue'))) :
                if IO[i].select('span.occurDate')[j].text == '合計輸入量' :
                    IO_list[i]['IO'][0] = 'I: '+IO[i].select('span.nValue')[j].text
                elif IO[i].select('span.occurDate')[j].text == '合計輸出量' :
                    IO_list[i]['IO'][1] = 'O: '+IO[i].select('span.nValue')[j].text
                elif IO[i].select('span.occurDate')[j].text == '差異量' :
                    IO_list[i]['IO'][2] = 'sum: '+IO[i].select('span.nValue')[j].text
        for k in range(len(IO_I[i].select('div.ioValue'))) :
            IO_list[i]['IO-I'].append(IO_I[i].select('div.ioValue')[k].text)
        for k in range(len(IO_O[i].select('div.ioValue'))) :
            IO_list[i]['IO-O'].append(IO_O[i].select('div.ioValue')[k].text)
    return IO_list
# ------------------------------
def GetLISList(ID, CHARTNO, sessions, lab = 'G') :
    """LIS List"""
    # ----檢查組合清單----
    lab_group_general = {
        'CBCDC' : { 'group' : ['WBC', 'RBC', 'HGB', 'HCT', 'MCV', 'MCH', 'MCHC', 'PLT', 'NEUT', 'LYMP', 'MONO', 'EOSIN', 'BASO', 'IG'], 'print' : ['WBC', 'HGB', 'PLT', 'NEUT', 'LYMP'], 'name' : 'B/R' },
        'E' : { 'group' : ['NA', 'KS', 'BUN', 'CREA', 'AST', 'ALT', 'CLS', 'GLUER'], 'print' : ['NA', 'KS', 'BUN', 'CREA', 'AST', 'ALT', 'CLS', 'GLUER'], 'name' : 'E' },
        'Protein' : { 'group' : ['TP', 'ALB', 'AG'], 'print' : ['TP', 'ALB', 'AG'], 'name' : 'TP/ALB/AG' },
        'CaMgP' : { 'group' : ['CA', 'CA1', 'IP', 'MG'], 'print' : ['CA', 'CA1', 'MG', 'IP'], 'name' : 'T./F.Ca/Mg/IP' },
        'ABG' : { 'group' : ['PHBG', 'PCO2', 'PO2', 'HCO3', 'TCO2', 'BE', 'SBE', 'SAT', 'SBC', 'O2CT'], 'print' : ['PHBG', 'PCO2', 'HCO3', 'SBE', 'PO2', 'SAT'], 'name' : 'ABG' }, 
        'CK/TnI' : { 'group' : ['CK', 'TROPI'], 'print' : ['CK', 'TROPI'], 'name' : 'CK/TnI' },
    }
    lab_group_ICU = {
        'CBC' : { 'group' : ['WBC', 'RBC', 'HGB', 'HCT', 'MCV', 'MCH', 'MCHC', 'PLT'], 'print' : ['WBC', 'HGB', 'PLT'], 'name' : 'WBC/Hgb/PLT' },
        'DC' : { 'group' : ['NEUT', 'LYMP', 'MONO', 'EOSIN', 'BASO', 'IG'], 'print' : ['NEUT', 'LYMP'], 'name' : 'NL' },
        'CRP/PCT/Lac' : { 'group' : ['CRP', 'PCT', 'LACT'], 'print' : ['CRP', 'PCT', 'LACT'], 'name' : 'CRP/PCT/Lac' },
        'Na/K/Cl' : { 'group' : ['NA', 'KS', 'CLS'], 'print' : ['NA', 'KS', 'CLS'], 'name' : 'Na/K/Cl' },
        'BUN/Cr' : { 'group' : ['BUN', 'CREA'], 'print' : ['BUN', 'CREA'], 'name' : 'BUN/Cr' },
        'CaMg' : { 'group' : ['CA', 'CA1', 'MG'], 'print' : ['CA', 'CA1', 'MG'], 'name' : 'T./F.Ca/Mg' },
        'AST/ALT' : { 'group' : ['AST', 'ALT'], 'print' : ['AST', 'ALT'], 'name' : 'AST/ALT' },
        'T./D.bil/TP' : { 'group' : ['TB', 'DBIL', 'TP'], 'print' : ['TB', 'DBIL', 'TP'], 'name' : 'T./D.bil/TP' },
        'CK/CKMB/TnI' : { 'group' : ['CK','CKMB', 'TROPI'], 'print' : ['CK','CKMB', 'TROPI'], 'name' : 'CK/CKMB/TnI' },
        'Amy/Lip/NH3' : { 'group' : ['AMY', 'LIPAS', 'NH3'], 'print' : ['AMY', 'LIPAS', 'NH3'], 'name' : 'Amy/Lip/NH3' },
        'BNP/D-dimer' : { 'group' : ['BNP', 'DDI'], 'print' : ['BNP', 'DDI'], 'name' : 'BNP/D-dimer' },
        'IP/Alk-P/r-GT' : { 'group' : ['IP', 'ALP', 'GGT'], 'print' : ['IP', 'ALP', 'GGT'], 'name' : 'IP/Alk-P/r-GT' },
        'BKet' : { 'group' : ['BKET'], 'print' : ['BKET'], 'name' : 'BKET' },
        'ALB' : { 'group' : ['ALB'], 'print' : ['ALB'], 'name' : 'ALB' },
    }
    lab_group_common = {
        'APTT PT' : { 'group' : ['PTP', 'PTT', 'INR', 'APTTP', 'APTTT'], 'print' : ['PTP', 'INR', 'APTTP'], 'name' : 'PT/INR/APTT' },
        'HAPTT PT' : { 'group' : ['HPT', 'HPTQC', 'HINR', 'HAPTT', 'APQC'], 'print' : ['HPT', 'HINR', 'HAPTT'], 'name' : 'HPT/INR/APTT' },
        'Lipid/UA' : { 'group' : ['HDLC', 'LDLC', 'TCHO', 'TG', 'UA'], 'print' : ['LDLC', 'TCHO', 'TG', 'UA'], 'name' : 'LDL/TCHO/TG/UA' },
        'VBG' : { 'group' : ['VPHBG', 'VPCO2', 'VPO2', 'VHCO3', 'VTCO2', 'VBE', 'VSBE', 'VSAT', 'VSBC', 'VO2CT'], 'print' : ['VPHBG', 'VPCO2', 'VHCO3', 'VSBE'], 'name' : 'VBG' },
        'U/R-bio' : { 'group' : ['GLUC', 'PRO', 'BIL', 'URO', 'PH', 'OB', 'KET', 'NIT', 'WBCU'], 'print' : ['GLUC', 'PRO', 'BIL', 'URO', 'PH', 'OB', 'KET', 'NIT', 'WBCU'], 'name' : 'U-Glu/Pro/Bil/Uro/pH/OB/Ket/Nit/sWBC' },
        'U/R-oth' : { 'group' : ['CLARI', 'SG', 'COLOR'], 'print' : ['CLARI', 'SG', 'COLOR'], 'name' : 'U-Clari/SG/Col' },
        'U/R-sed' : { 'group' : ['RBCU', 'WBCS', 'EPITU', 'CAST', 'CRYST', 'BAC', 'YE', 'SPERM'], 'print' : ['RBCU', 'WBCS', 'EPITU', 'CAST', 'CRYST', 'BAC', 'YE', 'SPERM'], 'name' : 'U-RBC/WBC/Epi/Cast/Crys/Bac/YE/Sperm' },
        'U10' : { 'group' : ['TPU', 'CREAU', 'BUNU', 'NAU', 'KU', 'CLU', 'UAU', 'IPU', 'CAU', 'MGU'], 'print' : ['TPU', 'CREAU', 'BUNU', 'NAU', 'KU', 'CLU', 'UAU', 'IPU', 'CAU', 'MGU'], 
                  'name' : 'U10-TP/Cr/BUN/Na/K/Cl/UA/IP/Ca/Mg' },
        '24HU' : { 'group' : ['TPUC', 'TCREAU', 'TBUNU', 'TNAU', 'TKU', 'TCLU', 'TUAU', 'TIPU', 'TCAU', 'TMGU'], 'print' : ['TPUC', 'TCREAU', 'TBUNU', 'TNAU', 'TKU', 'TCLU', 'TUAU', 'TIPU', 'TCAU', 'TMGU'], 
                  'name' : '24HU-TP/Cr/BUN/Na/K/Cl/UA/IP/Ca/Mg' },
        'st/R' : { 'group' : ['APP', 'OV', 'GB', 'OCB', 'CO', 'RBCS', 'WBCT', 'MUC', 'FATS', 'OVCM'], 'print' : ['APP', 'OV', 'GB', 'OCB', 'CO', 'RBCS', 'WBCT', 'MUC', 'FATS'], 'name' : 'Stool' },
        'SK/L' : { 'group' : ['KAP', 'LAM', 'KLC'], 'print' : ['KAP', 'LAM', 'KLC'], 'name' : 'K/L chain' },
        'UK/L' : { 'group' : ['AMT10', 'IGU', 'KAPU', 'LAMU'], 'print' : ['KAPU', 'LAMU'], 'name' : 'U-K/L chain' },
        'IGX' : { 'group' : ['IGG', 'IGM', 'IGA', 'IGD', 'IGE'], 'print' : ['IGG', 'IGM', 'IGA', 'IGD', 'IGE'], 'name' : 'IgGMADE' },
        'C3/C4' : { 'group' : ['C3', 'C4'], 'print' : ['C3', 'C4'], 'name' : 'C3/4' },
        'CBCH' : { 'group' : ['HWBC', 'HRBC', 'HHGB', 'HHCT', 'HMCV', 'HMCH', 'HMCHC', 'HPLT', 'HRDWS', 'HRDWC', 'HPDW', 'HMPV', 'HPLCR', 'HPCT', 'SCHIS'], 
                  'print' : ['HWBC', 'HHGB', 'HMCV', 'HPLT', 'SCHIS'], 'name' : 'CBCH(WBC/Hg/MCV/PLT/schi)'},
        'DCH' : { 'group' : ['HATLY', 'HMYEL', 'HMETA', 'NUCLE', 'HNEUT', 'HLYMP', 'HMONO', 'HEOS', 'HBASO', 'COUNT', 'HWBC', 'HNE'],
                'print' : ['HATLY', 'HMYEL', 'HMETA', 'NUCLE', 'HNEUT', 'HLYMP', 'HMONO', 'HEOS', 'HBASO', 'HWBC', 'HNE'],
                'name' : 'CBCH(ATLY/MYEL/META/NUC/NEU/LYM/MON/EOS/BASO/WBC/NE)' },
        'anemia' : { 'group' : ['FE', 'FERRI', 'TIBC', 'Vitamin B12', 'FOLAT'], 'print' : ['FE', 'FERRI', 'TIBC', 'Vitamin B12', 'FOLAT'], 'name' : 'Fe/Ferri/TIBC/B12/Folate' },
        
    }
    # -------------------------
    if lab == 'I' : lab_group = {**lab_group_ICU, **lab_group_common}
    else : lab_group = {**lab_group_general, **lab_group_common}
    # -------------------------
    print('LIS', end = '->')
    lab_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/LisList.aspx?login_id='+ID+'&special=n&cno='+CHARTNO)
    soup = BeautifulSoup(lab_requ.text, 'html.parser')
    lab_d = soup.select('table#GridView2 td a')
    lab_t = soup.select('table#GridView2 td pre')
    
    lab_total = []
    culture = { 'access' : [], 'collect' : [], 'type' : [], 'org' : [], 'other' : [], }
    other = ['Stain report(AFB)', 'Aerobic result', 'Result BCU', 'RESULT TB CULTURE' ,'RESULT FUNG STAIN', 'RESULT FUNG CULTURE', 'Cryptococcus Ag']
    
    for i in range(min(3,len(lab_d))) :
        lab_data = {
        'date' : lab_d[i].text,
        'type' : lab_t[i].text.split('\r\n'),
        'name' : [],
        'value' : [],
        }
        lab_total.append(lab_data)

    for d in range(min(3,len(lab_total))) :
        data_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/LisREPORT.aspx?login_id='+ID+'&LASTUPDTESTDATE='+lab_total[d]['date'].replace('.','')+'&TITLE='+lab_total[d]['date']+'&special=n&cno='+CHARTNO)
        soup = BeautifulSoup(data_requ.text, 'html.parser')
        lab_data_soup = soup.select('table#GridView1 td')

        lab_temp = { 'name' : [], 'collect' : [], 'value' : [], 'abnl' : [], }
        for i in range(2,len(lab_data_soup),17) :
            n = lab_data_soup[i+14].text
            if sum([ 1 for t in n if t in '1234568790'])>1 : n = lab_data_soup[i].text
            lab_temp['name'].append(n.replace('YPLT', 'PLT').replace('PTL','PLT'))
            abnl = re.search(r'bgcolor=\"(.*)\"', str(lab_data_soup[i+1]))
            if abnl is not None : 
                abnl = abnl.group(1)
            if abnl == 'Pink' :
                ab = 'H'
            elif abnl == 'Yellow' :
                ab = 'L'
            else : ab =''
            lab_temp['collect'].append(lab_data_soup[i+6].text)
            lab_temp['value'].append(lab_data_soup[i+1].text+ab)
            lab_temp['abnl'].append(ab != '')
        collect_list = sorted(set(lab_temp['collect']), reverse = True)
        
        for j in range(len(collect_list)) :
            key_collect = [ x for x in range(len(lab_temp['collect'])) if lab_temp['collect'][x] == collect_list[j]]
            name_sel = [lab_temp['name'][x] for x in key_collect]
            data_sel = [lab_temp['value'][x] for x in key_collect]
            abnl_sel = [lab_temp['abnl'][x] for x in key_collect]
            # ---group formation---
            group = list(lab_group)
            for g in group :
                if any([ name in name_sel for name in lab_group[g]['group'] ]) :
                    lab_total[d]['name'].append(lab_group[g]['name'])
                    temp_v = ''
                    for p in lab_group[g]['print'] :
                        if p in name_sel : temp_v = temp_v+data_sel[name_sel.index(p)]+'/'
                        else : temp_v = temp_v+'-/'
                    lab_total[d]['value'].append(temp_v.rstrip('/'))
                    key_collect = [ k for k in key_collect if lab_temp['name'][k] not in lab_group[g]['group'] ]
            # ---residue---
            for k in key_collect :
                lab_total[d]['name'].append(lab_temp['name'][k])
                lab_total[d]['value'].append(lab_temp['value'][k])
            
        # ---culture result---
        culture_soup = soup.select('table#GridView2')
        if culture_soup != [] :
            culture_soup = culture_soup[0].select('td')
            cul_raw = {
                'name' : [],
                'value' : [],
                'collect' : [],
                'access' : [],
            }
            for c in range(2,len(culture_soup),14)  :
                cul_raw['name'].append(culture_soup[c].text)
                cul_raw['value'].append(culture_soup[c+1].text)
                cul_raw['collect'].append(culture_soup[c+3].text)
                cul_raw['access'].append(culture_soup[c+8].text)

            access_list = sorted(set(cul_raw['access']), reverse = True)
            for a in access_list :
                if a not in culture['access'] :
                    culture['access'].append(a)
                    culture['collect'].append('')
                    culture['type'].append('')
                    culture['org'].append([])
                    culture['other'].append([])

            for a in access_list :
                cul_index = culture['access'].index(a)
                key_access = [ x for x in range(len(cul_raw['access'])) if cul_raw['access'][x] == a]
                name_sel = [ cul_raw['name'][x] for x in key_access ]
                value_sel = [ cul_raw['value'][x] for x in key_access ]
                culture['collect'][cul_index] = [ cul_raw['collect'][x][4:8] for x in key_access ][0]
                if 'Sample type' in name_sel :
                    culture['type'][cul_index] = value_sel[name_sel.index('Sample type')]
                for n in range(len(name_sel)) :
                    if 'Organism' in name_sel[n] :
                        if value_sel[n] not in culture['org'][cul_index] : culture['org'][cul_index].append(value_sel[n])
                    if name_sel[n] in other :
                        if value_sel[n] not in culture['other'][cul_index] : culture['other'][cul_index].append(name_sel[n]+':'+value_sel[n])
    return [lab_total, culture]
# ------------------------------
def GetRISList(ID, CHARTNO, sessions) :
    """RIS List"""
    print('RIS', end = '->')
    RIS_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/RisList.aspx?login_id='+ID+'&special=n&cno='+CHARTNO)
    soup = BeautifulSoup(RIS_requ.text, 'html.parser')
    RIS = soup.select('table#GridView1 td')
    RIS_list = []
    CXR_url = {
        'date' : [],
        'url' : [],
    }
    for R in range(1,len(RIS),21) :
        if RIS[R].text.replace('\n','') != "Echo(-M-mode & -sector- scan)" :
            txt = RIS[R].text.replace('\n','')
            txt = re.sub(r".*CHEST.*-[^C].*","CXR",txt).replace('CHEST  WITHOUT/WITH CONTRAST-C.T.','CHE CNYCT').replace('K.U.B. (SUPINE)','KUB').replace(', 2 VIEWS (A-P & LAT.)','').replace('SKULL ROUTINE (P-A & LAT.)','SKULL')
            txt = txt.replace('Tl-201 myocardial perfusion scan','Tl201').replace('SONO. WHOLE ABDOMEN STUDY','AbdUS').replace('Echo-Doppler echo','CVecho')
            txt = txt.replace('Upper GI panendoscopy','EGD').replace('Whole body inflammation scan','InfScan').replace('WITHOUT CONTRAST-C.T.','HRCT').replace('ABDOMEN  C.T. without contrast','AbdCNCT')
            txt = txt.replace('ABDOMEN  WITHOUT/WITH CONTRAST-C.T.','AbdCNYCT').replace('91 SCREENING-MAMMOGRAPHY','MAMMO')
            txt = re.sub(r',.*\'T','',txt).replace('X-ray bone densitometry','DEXA')
            if txt == 'CXR' :
                CXR = sessions.get('http://html5pacs.ndmctsgh.edu.tw/DicomWeb/DicomWeb.dll/WADO?requestType=Query&User=ebm&Password=pacs&lineNumberPerPage=25&displayLevel=image&noImageRec=1&incSeriesThumb=1&shareID=&sharePassword=&patientID='+CHARTNO+'&patientName=&studyDate=&accessionNumber='+RIS[R+7].text+'&studyID=&modality=&bodyPart=&pageNumber=undefined')
                if re.search(r'thumbnailURL="(.*)"', CXR.text) is not None :
                    url = 'http://html5pacs.ndmctsgh.edu.tw'+re.search(r'thumbnailURL="(.*)"', CXR.text).group(1)
                    CXR_url['url'].append(url)
                    CXR_url['date'].append(RIS[R+1].text)
            RIS_list.append(RIS[R+1].text+'-'+txt)
    return [RIS_list, CXR_url]
# ------------------------------
def GetImpList(ID, CHARTNO, INDATETIME, HCASENO, sessions) :
    """Impression List"""
    print('Dx', end = '->')
    imp_list = { 'ad' : [], 'addate' : '', 'EROPD' : [], 'EROPDdate' : '', 'dc' : [], 'dcdate' : '' }
    # ---AD---
    ad_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/AdmissionnoteLIST.aspx?login_id='+ID+'&special=n&cno='+PtList['CHARTNO'][pt])
    soup = BeautifulSoup(ad_requ.text, 'html.parser')
    ad = soup.select('table#GridView1 td')
    if ad != [] :
        if ad[5].text == PtList['INDATETIME'][pt][0:7] :
            note_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/AdmissionnoteREPORT.aspx?login_id='+ID+'&HCASENO='+PtList['HCASENO'][pt]+'&TITLE='+ad[5].text+'-'+ad[10].text+'&special=n&cno='+PtList['CHARTNO'][pt])
            note_txt = note_requ.text.replace('\n','')
            if re.search(r'\$(.*?)-*\*Diagnostic',note_txt) is not None :
                imp_list['ad'] = re.search(r'\$(.*?)-*\*Diagnostic',note_txt).group(1).split('*.')[1::]
                imp_list['addate'] = ad[5].text
    # ---EROPD---
    ad_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/OPDList.aspx?login_id='+ID+'&special=n&cno='+PtList['CHARTNO'][pt])
    soup = BeautifulSoup(ad_requ.text, 'html.parser')
    ad = soup.select('table#GridView1 td')
    if ad != [] :
        if '急診' in ad[0].text or PtList['DepName'][pt] in ad[9].text : 
            note_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/OPDREPORT.aspx?login_id='+ID+'&RID='+ad[12].text+'&TITLE='+ad[0].text+'&special=n&diagdate='+ad[7].text+'&cno='+PtList['CHARTNO'][pt])
            note_requ = note_requ.text.replace('<br/>','').replace('\n','')
            imp_list['EROPDdate'] = ad[7].text
            imp_list['EROPD'] = re.search(r'初步診斷.*:(.*)主診斷', note_requ)
            if imp_list['EROPD'] is not None : imp_list['EROPD'] = imp_list['EROPD'].group(1).split('*')[1::]
            else : imp_list['EROPD'] = [re.search('診斷名稱: (.*?)處方', note_requ).group(1)]
    # ---DC---
    ad_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/DischargenoteList.aspx?login_id='+ID+'&special=n&cno='+PtList['CHARTNO'][pt])
    soup = BeautifulSoup(ad_requ.text, 'html.parser')
    ad = soup.select('table#GridView1 td')
    if ad != [] :
        if ad[12].text == PtList['DepName'][pt] :
            note_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/DischargenoteREPORT.aspx?login_id='+ID+'&HCASENO='+ad[15].text+'&TITLE='+ad[6].text+'-'+ad[11].text+'-'+ad[12].text+'&special=n&cno='+PtList['CHARTNO'][pt])
            note_txt = note_requ.text.replace('-','').split('\n')
            note_txt = [ n for n in note_txt if n != '']
            imp_list['dcdate'] = ad[6].text
            imp_list['dc'] = [ n for n in note_txt[note_txt.index('*出院診斷(Discharge Diagnosis):')+1:note_txt.index('*主訴(Chief ComPlaint):')] if '*.' in n ]
    imp = []
    if imp_list['ad'] != [] : imp += imp_list['ad']
    elif (imp_list['EROPDdate'] > imp_list['dcdate']) and (imp_list['EROPD'] != []) : imp += imp_list['EROPD']
    if imp_list['dc'] != [] : imp += ['上次出院診斷']+imp_list['dc']
    if imp == [] : imp = ['目前無登錄診斷']
    for i in range(len(imp)) : imp[i] = '>'+imp[i].replace('*.','')
    return imp
# ------------------------------
def GetMedList(ID, CHARTNO, sessions) :
    """Med List"""
    print('med', end = '->')
    med_requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/mr/HISEXNDREPORT.aspx?login_id='+ID+'&special=n&cno='+CHARTNO)
    soup = BeautifulSoup(med_requ.text, 'html.parser')
    if soup.select('pre') != [] :
        med_list = soup.select('pre')[0].text.replace('[有效藥囑]\n\n處     方                                 劑量  頻率    用法  開始時間     結束時間     \n','')
    else : med_list = ''
    # ---繪圖---
    font = ImageFont.truetype('mingliu.ttc', 9)
    ascent, descent = font.getmetrics()
    img = Image.new('RGB', (10,10), color = (255, 255, 255))
    d = ImageDraw.Draw(img)
    txt_s = d.textsize(med_list, font, spacing = 1)
    txt_h = txt_s[1]+descent
    txt_w = txt_s[0]+1
    img = img.resize((txt_w, txt_h))
    d = ImageDraw.Draw(img)
    d.text((0,0), med_list, font = font, spacing = 1, fill=(0,0,0))
    output_buffer = BytesIO()
    img.save(output_buffer, format='JPEG')
    return output_buffer
# ------------------------------
def GetICUsheet(CHARTNO, sessions) :
    """ICU flowsheet"""
    print('ICU flowsheet', end = '->')
    MI_chart = {
        'T' : '', 'P' : '', 'R' : '',
        'SBP' : '', 'DBP' : '', 'MAP' : '',
        'VenMode' : '', 'FiO2' : '', 'rate' : '', 'Pressure' : '', 'PEEP' : '',
        'GCS-E' : '', 'GCS-M' : '', 'GCS-V' : '',
        'pH' : '', 'PCO2' : '', 'HCO3' : '', 'PO2' : '', 'SaO2' : '',
        'IO' : { 'date' : '', 'IV' : '', 'CPN/PPN' : '', 'transfusion' : '', 'GI' : '', 'IT' : '', 'UO' : '', 'stool' : '', 'drain' : '', 'vomit' : '', 'DeHy' : '', 'OT' : '', 'IO' : ''},
        }
    
    ICU_requ = sessions.get('http://ivue.ndmctsgh.edu.tw/iVue/patient.aspx?ChartNo='+CHARTNO)
    soup = BeautifulSoup(ICU_requ.text, 'html.parser')
    k = soup.select('td.normC a ')
    if k != [] :
        ICCA_url = []
        for i in k[:12] : 
            ICCA_url.append('http://ivue.ndmctsgh.edu.tw/iVue/'+i['href'])
        s = sessions.get(ICCA_url[2])
        soup = BeautifulSoup(s.text, 'html.parser').select('table.mainTBL td')
        lines = BeautifulSoup(s.text, 'html.parser').select('table.mainTBL tr')
        list_len = int(len(soup)/len(lines))
        ICU_flowsheet = []
        for i in range(0,len(soup),list_len) :
            ICU_flowsheet.append({ soup[i].text : [ soup[x].text for x in range(i+1,i+12) if soup[x].text != '' ] })
        for ICU in ICU_flowsheet :
            if list(ICU)[0] == '睜眼反應' and ICU[list(ICU)[0]] != [] : 
                MI_chart['GCS-E'] = 'E'+ICU[list(ICU)[0]][-1].split('-')[-1]
            if list(ICU)[0] == '最佳運動反應' and ICU[list(ICU)[0]] != [] : 
                MI_chart['GCS-M'] = 'M'+ICU[list(ICU)[0]][-1].split('-')[-1]
            if list(ICU)[0] == '最佳言辭反應' and ICU[list(ICU)[0]] != [] : 
                MI_chart['GCS-V'] = 'V'+ICU[list(ICU)[0]][-1].split('-')[-1]
            if list(ICU)[0] == '體溫℃' and ICU[list(ICU)[0]] != [] : 
                MI_chart['T'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == '脈搏/分' and ICU[list(ICU)[0]] != [] : 
                MI_chart['P'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == '呼吸/分' and ICU[list(ICU)[0]] != [] : 
                MI_chart['R'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == ' \xa0\xa0 - Systolic' and ICU[list(ICU)[0]] != [] : 
                MI_chart['SBP'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == ' \xa0\xa0 - Diastolic' and ICU[list(ICU)[0]] != [] : 
                MI_chart['DBP'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == ' \xa0\xa0 - Mean' and ICU[list(ICU)[0]] != [] : 
                MI_chart['MAP'] = ICU[list(ICU)[0]][-1].split(' ')[0]
        s = sessions.get(ICCA_url[8])
        soup = BeautifulSoup(s.text, 'html.parser').select('table.mainTBL td')
        lines = BeautifulSoup(s.text, 'html.parser').select('table.mainTBL tr')
        list_len = int(len(soup)/len(lines))
        ICU_flowsheet = []
        for i in range(0,len(soup),list_len) :
            ICU_flowsheet.append({ soup[i].text : [ soup[x].text for x in range(i+1,i+12) if soup[x].text != '' ] })
        for ICU in ICU_flowsheet :
            if list(ICU)[0] == 'Vent Mode' and ICU[list(ICU)[0]] != [] : 
                MI_chart['VenMode'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'Vent Rate-Total' and ICU[list(ICU)[0]] != [] : 
                MI_chart['rate'] = ICU[list(ICU)[0]][-1].split(' ')[-1]
            if list(ICU)[0] == 'FiO2 (%)' and ICU[list(ICU)[0]] != [] : 
                MI_chart['FiO2'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'PS Level' and ICU[list(ICU)[0]] != [] : 
                MI_chart['Pressure'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == 'PEEP' and ICU[list(ICU)[0]] != [] : 
                MI_chart['PEEP'] = ICU[list(ICU)[0]][-1].split(' ')[0]
            if list(ICU)[0] == 'pH' and ICU[list(ICU)[0]] != [] : 
                MI_chart['pH'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'PaCO2' and ICU[list(ICU)[0]] != [] : 
                MI_chart['PCO2'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'HCO3' and ICU[list(ICU)[0]] != [] : 
                MI_chart['HCO3'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'PaO2' and ICU[list(ICU)[0]] != [] : 
                MI_chart['PO2'] = ICU[list(ICU)[0]][-1]
            if list(ICU)[0] == 'SaO2' and ICU[list(ICU)[0]] != [] : 
                MI_chart['SaO2'] = ICU[list(ICU)[0]][-1]
        s = sessions.get(ICCA_url[11])
        soup = BeautifulSoup(s.text, 'html.parser').find_all(onclick = 'onClickIODate(this)')
        io_url = [ i.a['href'] for i in soup if i.a != None ]
        IOtable = { 'date' : '', 'IV' : '', 'CPN/PPN' : '', 'transfusion' : '', 'GI' : '', 'IT' : '', 'UO' : '', 'stool' : '', 'drain' : '', 'vomit' : '', 'DeHy' : '', 'OT' : '', 'IO' : ''}
        sDate = datetime.date.today() - datetime.timedelta(1)
        sDate = sDate.strftime('%Y-%m-%d')
        yesday = [ url for url in io_url if sDate in url]
        if yesday != [] :
            s = sessions.get(yesday[0])
            IOtable['date'] = re.search(r'IODate=(.*)',yesday[0]).group(1)
            soup = BeautifulSoup(s.text, 'html.parser')
            for i in range(len(list(IOtable))-1) : IOtable[list(IOtable)[i+1]] = soup.select('tr')[-1].select('td')[1+i].text
        elif io_url != []: 
            s = sessions.get(io_url[0])
            IOtable['date'] = re.search(r'IODate=(.*)',io_url[0]).group(1)
            soup = BeautifulSoup(s.text, 'html.parser')
            for i in range(len(list(IOtable))-1) : IOtable[list(IOtable)[i+1]] = soup.select('tr')[-1].select('td')[1+i].text
        MI_chart['IO'] = IOtable  
    return MI_chart
# ---詳細格式---
def detailform(PtList, Ptnum) :
    doc = Document('default.docx')
    for pt in range(Ptnum) :
        print('Pt[',pt+1,']->', end="")
        section = doc.sections[-1]
        style = doc.styles['Normal']
        style.font.size = Pt(9)
        # ---版面配置---
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        # ---新增ptlist表格---
        # ---General data---
        table1 = doc.add_table(rows = 41, cols = 9, style = 'Table Grid')
        table1.allow_autofit = True
        table1.cell(0,0).text = PtList['NRCODE'][pt]+'-'+PtList['BEDNO'][pt]
        table1.cell(0,1).merge(table1.cell(0,2))
        table1.cell(0,1).text = PtList['NameGenderAge'][pt]
        table1.cell(0,3).text = PtList['CHARTNO'][pt]
        table1.cell(0,4).text = '入院'
        table1.cell(0,5).merge(table1.cell(0,6))
        table1.cell(0,5).text = PtList['INDATETIME'][pt][:7]
        table1.cell(0,7).merge(table1.cell(0,8))
        table1.cell(0,7).text = 'MBD□'
        table1.cell(1,0).merge(table1.cell(8,3))
        table1.cell(1,0).text = '\n'.join(PtList['PtData'][pt]['imp'])
        # ---用藥轉圖片---
        table1.cell(1,4).merge(table1.cell(8,8))
        m = table1.cell(1,4).add_paragraph().add_run()
        m.add_picture(PtList['PtData'][pt]['med'])
        # ---TPR chart---
        for i in range(3) :
            table1.cell(9,i*3).merge(table1.cell(9,i*3+2))
            table1.cell(10,i*3).merge(table1.cell(15,i*3+1))
            table1.cell(16,i*3).merge(table1.cell(21,i*3+1))
            table1.cell(11,i*3+2).merge(table1.cell(17,i*3+2))
        for i in range(min(3,len(PtList['PtData'][pt]['TPR']))) :
            table1.cell(9,i*3).text = PtList['PtData'][pt]['TPR'][i]['date']
            table1.cell(10,i*3).text = '\n'.join(PtList['PtData'][pt]['TPR'][i]['TPR'])
            table1.cell(16,i*3).text = '\n'.join(PtList['PtData'][pt]['IO'][i]['BP'])
            table1.cell(11,i*3+2).text = '\n'.join(PtList['PtData'][pt]['IO'][i]['SpO2'])
            table1.cell(18,i*3+2).text = 'I/O'
            for j in range(3) : table1.cell(19+j,i*3+2).text = PtList['PtData'][pt]['IO'][i]['IO'][j]
        # ---Lab data---
        table1.cell(22,0).merge(table1.cell(22,5))
        table1.cell(22,0).text = 'Lab Data'
        table1.cell(22,0).alignment = 1
        for i in range(3) :
            table1.cell(23,i*2).merge(table1.cell(23,i*2+1))
            table1.cell(24,i*2).merge(table1.cell(34,i*2+1))
        l = min(3,len(PtList['PtData'][pt]['LIS']))-1
        for i in range(l,-1,-1) :
            table1.cell(23,(l-i)*2).text = PtList['PtData'][pt]['LIS'][i]['date']
            data = [ PtList['PtData'][pt]['LIS'][i]['name'][x]+'-'+PtList['PtData'][pt]['LIS'][i]['value'][x] for x in range(len(PtList['PtData'][pt]['LIS'][i]['name'])) ]
            table1.cell(24,(l-i)*2).text = '\n'.join(data)
        # ---Culture---
        table1.cell(22,6).merge(table1.cell(22,8))
        table1.cell(22,6).text = 'Infection'
        table1.cell(23,6).merge(table1.cell(25,8))
        INF = []
        for i in range(len(PtList['PtData'][pt]['culture']['access'])) :
            inf_temp = PtList['PtData'][pt]['culture']['collect'][i]+PtList['PtData'][pt]['culture']['type'][i]+'-'
            if PtList['PtData'][pt]['culture']['org'][i] != [] : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['org'][i])
            else : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['other'][i])
            INF.append(inf_temp)
        table1.cell(23,6).text = '\n'.join(INF)
        # ---Anti---
        table1.cell(26,6).merge(table1.cell(26,8))
        table1.cell(26,6).text = 'Anti'
        table1.cell(27,6).merge(table1.cell(29,8))
        anti = [ PtList['PtData'][pt]['anti']['name'][x]+'-'+PtList['PtData'][pt]['anti']['dose'][x]+'-'+PtList['PtData'][pt]['anti']['date'][x] for x in range(len(PtList['PtData'][pt]['anti']['name'])) ]
        table1.cell(27,6).text = '\n'.join(anti)
        # ---Image---
        table1.cell(35,0).merge(table1.cell(35,4))
        table1.cell(35,0).text = 'Image'
        table1.cell(36,0).merge(table1.cell(40,4))
        table1.cell(36,0).text = '\n'.join(PtList['PtData'][pt]['RIS'][:min(5,len(PtList['PtData'][pt]['RIS']))-1])

        # ---CXR近期兩張---
        for i in range(2) :
            table1.cell(35,i*2+5).merge(table1.cell(35,i*2+6))
            table1.cell(36,i*2+5).merge(table1.cell(40,i*2+6))
        for i in range(min(2,len(PtList['PtData'][pt]['CXR']['date']))) :
            table1.cell(35,i*2+5).text = PtList['PtData'][pt]['CXR']['date'][i]
            img = sessions.get(PtList['PtData'][pt]['CXR']['url'][i]).content
            r = table1.cell(36,i*2+5).add_paragraph().add_run()
            r.add_picture(BytesIO(img))

        # ---note欄---
        table1.cell(30,6).merge(table1.cell(30,8))
        table1.cell(30,6).text = 'note'
        table1.cell(31,6).merge(table1.cell(34,8))
        doc.add_section()
    print(doc)
    return doc

# ---簡易格式---
def simpleform(PtList, Ptnum) :
    doc = Document('default.docx')
    for pt in range(Ptnum) :
        print('Pt[',pt+1,']->', end="")
        section = doc.sections[-1]
        style = doc.styles['Normal']
        style.font.size = Pt(9)
        # ---版面配置---
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        # ---新增ptlist表格---
        # ---General data---
        table1 = doc.add_table(rows = 2, cols = 6, style = 'Table Grid')
        table1.allow_autofit = False
        table1.cell(0,0).merge(table1.cell(1,0))
        table1.cell(0,0).text = PtList['CHARTNO'][pt]+'\n'+PtList['NRCODE'][pt]+'-'+PtList['BEDNO'][pt]+'\n'+PtList['NameGenderAge'][pt]+'\n'+PtList['DepName'][pt]+'\nVS: '+PtList['VSDRNAME'][pt]+'\n'+PtList['INDATETIME'][pt][:7]
        table1.cell(0,1).merge(table1.cell(1,1))
        table1.cell(0,1).text = '\n'.join(PtList['PtData'][pt]['imp'])
        # ---用藥轉圖片---
        table1.cell(1,3).merge(table1.cell(1,4))
        m = table1.cell(1,3).add_paragraph().add_run()
        m.add_picture(PtList['PtData'][pt]['med'])
        # ---TPR chart---
        TPR = []
        BP = []
        d = range(-1*min(2,len(PtList['PtData'][pt]['TPR'])),0)
        for i in d :
            TPR = TPR+[PtList['PtData'][pt]['TPR'][i]['date']]+PtList['PtData'][pt]['TPR'][i]['TPR']
            BP = BP+[PtList['PtData'][pt]['IO'][i]['date']]+PtList['PtData'][pt]['IO'][i]['BP']
        table1.cell(0,3).text = '\n'.join(TPR)
        table1.cell(0,4).text = '\n'.join(BP)
        table1.cell(0,5).merge(table1.cell(1,5))
        table1.cell(0,5).text = '\n'.join(PtList['PtData'][pt]['IO'][-1]['IO'])
        # ---Lab data---
        table1.cell(0,2).merge(table1.cell(1,2))
        l = min(2,len(PtList['PtData'][pt]['LIS']))
        data = []
        for i in range(l) :
            table1.cell(0,2).add_paragraph(PtList['PtData'][pt]['LIS'][i]['date'])
            data = [ PtList['PtData'][pt]['LIS'][i]['name'][x]+'-'+PtList['PtData'][pt]['LIS'][i]['value'][x] for x in range(len(PtList['PtData'][pt]['LIS'][i]['name'])) ]
            table1.cell(0,2).add_paragraph('\n'.join(data))
        # ---Culture---
        INF = []
        for i in range(len(PtList['PtData'][pt]['culture']['access'])) :
            inf_temp = PtList['PtData'][pt]['culture']['collect'][i]+PtList['PtData'][pt]['culture']['type'][i]+'-'
            if PtList['PtData'][pt]['culture']['org'][i] != [] : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['org'][i])
            else : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['other'][i])
            INF.append(inf_temp)
        table1.cell(0,2).add_paragraph('\n'.join(INF))
        # ---Anti---
        anti = [ PtList['PtData'][pt]['anti']['name'][x]+'-'+PtList['PtData'][pt]['anti']['dose'][x]+'-'+PtList['PtData'][pt]['anti']['date'][x] for x in range(len(PtList['PtData'][pt]['anti']['name'])) ]
        table1.cell(1,3).add_paragraph('\n'.join(anti))
        # ---Image---
        table1.cell(0,5).add_paragraph('\n'.join(PtList['PtData'][pt]['RIS'][:min(3,len(PtList['PtData'][pt]['RIS']))-1]))
        # ---CXR近期一張---
        for i in range(min(1,len(PtList['PtData'][pt]['CXR']['date']))) :
            table1.cell(0,5).add_paragraph(PtList['PtData'][pt]['CXR']['date'][i])
            img = sessions.get(PtList['PtData'][pt]['CXR']['url'][i]).content
            r = table1.cell(0,5).add_paragraph().add_run()
            r.add_picture(BytesIO(img))
    print(doc)
    return doc
# ---MI chart update格式---
def MIform(PtList, Ptnum) :
    doc = Document('default.docx')
    for pt in range(Ptnum) :
        print('Pt[',pt+1,']->', end="")
        section = doc.sections[-1]
        style = doc.styles['Normal']
        style.font.size = Pt(9)
        # ---版面配置---
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        # ---新增ptlist表格---
        # ---General data---
        table1 = doc.add_table(rows = 41, cols = 5, style = 'Table Grid')
        table1.allow_autofit = False
        table1.cell(0,0).text = PtList['BEDNO'][pt]
        table1.cell(0,1).text = PtList['NameGenderAge'][pt]
        table1.cell(0,2).text = PtList['CHARTNO'][pt]
        table1.cell(0,3).text = PtList['SECTIONNO'][pt]
        table1.cell(0,4).text = PtList['INDATETIME'][pt][3:7]
        table1.cell(1,0).merge(table1.cell(1,4))
        table1.cell(1,0).text = '\n'.join(PtList['PtData'][pt]['imp'])
        # ---用藥轉圖片---
        table1.cell(3,3).merge(table1.cell(40,4))
        m = table1.cell(3,3).add_paragraph().add_run()
        m.add_picture(PtList['PtData'][pt]['med'])
        # ---Anti---
        anti = [ PtList['PtData'][pt]['anti']['name'][x]+'-'+PtList['PtData'][pt]['anti']['dose'][x]+'-'+PtList['PtData'][pt]['anti']['date'][x] for x in range(len(PtList['PtData'][pt]['anti']['name'])) ]
        table1.cell(3,3).add_paragraph('\n'.join(anti))
        # ---date---
        table1.cell(2,0).text = 'Date'
        table1.cell(2,1).text = datetime.date.today().strftime('%m%d')
        # ---Sedation---
        table1.cell(3,0).text = 'Sedation'
        table1.cell(3,1).text = PtList['PtData'][pt]['MIchart']['GCS-E']+'/'+PtList['PtData'][pt]['MIchart']['GCS-M']+'/'+PtList['PtData'][pt]['MIchart']['GCS-V']
        # ---Perfusion---
        table1.cell(6,0).text = 'Perfusion'
        table1.cell(7,0).text = 'T/P/R'
        table1.cell(7,1).text = PtList['PtData'][pt]['MIchart']['T']+'/'+PtList['PtData'][pt]['MIchart']['P']+'/'+PtList['PtData'][pt]['MIchart']['R']
        table1.cell(8,0).text = 'BP'
        table1.cell(8,1).text = PtList['PtData'][pt]['MIchart']['SBP']+'/'+PtList['PtData'][pt]['MIchart']['DBP']+'/'+PtList['PtData'][pt]['MIchart']['MAP']
        # ---Ventilation---
        table1.cell(10,0).text = 'Ventilation(CXR)'
        table1.cell(11,0).text = 'Mode'
        table1.cell(11,1).text = PtList['PtData'][pt]['MIchart']['VenMode']
        table1.cell(12,0).text = 'FiO2/rate'
        table1.cell(12,1).text = PtList['PtData'][pt]['MIchart']['FiO2']+'/'+PtList['PtData'][pt]['MIchart']['rate']
        table1.cell(13,0).text = 'Pressure/PEEP'
        table1.cell(13,1).text = PtList['PtData'][pt]['MIchart']['Pressure']+'/'+PtList['PtData'][pt]['MIchart']['PEEP']
        table1.cell(14,0).text = 'pH/pCO2/HCO3'
        table1.cell(14,1).text = PtList['PtData'][pt]['MIchart']['pH']+'/'+PtList['PtData'][pt]['MIchart']['PCO2']+'/'+PtList['PtData'][pt]['MIchart']['HCO3']
        table1.cell(15,0).text = 'PO2/SaO2'
        table1.cell(15,1).text = PtList['PtData'][pt]['MIchart']['PO2']+'/'+PtList['PtData'][pt]['MIchart']['SaO2']
        # ---Nutrition---
        table1.cell(16,0).text = 'Nutrition'
        table1.cell(17,0).text = 'I/O'
        table1.cell(17,1).text = PtList['PtData'][pt]['MIchart']['IO']['IO']
        table1.cell(18,0).text = 'IV/輸血'
        table1.cell(18,1).text = PtList['PtData'][pt]['MIchart']['IO']['IV']+'/'+PtList['PtData'][pt]['MIchart']['IO']['transfusion']
        table1.cell(19,0).text = 'PO/PPN or CPN'
        table1.cell(19,1).text = PtList['PtData'][pt]['MIchart']['IO']['GI']+'/'+PtList['PtData'][pt]['MIchart']['IO']['CPN/PPN']
        table1.cell(20,0).text = 'Urine/Stool'
        table1.cell(20,1).text = PtList['PtData'][pt]['MIchart']['IO']['UO']+'/'+PtList['PtData'][pt]['MIchart']['IO']['stool']
        table1.cell(21,0).text = 'Drain'
        table1.cell(21,1).text = PtList['PtData'][pt]['MIchart']['IO']['drain']

        lab_item = ['ALB','WBC/Hgb/PLT','N/L','CRP/PCT/Lac','Na/K/Cl','BUN/Cr','T./F.Ca/Mg','AST/ALT','PT/INR/PTT','T./D.bil/TP','CK/CKMB/TnI','Amy/Lip/NH3','Lipid/UA','BNP/D-dimer','IP/Alk-P/r-GT','BKET']
        if PtList['PtData'][pt]['LIS'][0]['date'] == datetime.date.today().strftime('%Y.%m.%d') :
            if lab_item[0] in PtList['PtData'][pt]['LIS'][0]['name'] : table1.cell(22,1).text = PtList['PtData'][pt]['LIS'][0]['value'][PtList['PtData'][pt]['LIS'][0]['name'].index(lab_item[0])] 
            for i in range(1,16) :
                if lab_item[i] in PtList['PtData'][pt]['LIS'][0]['name'] : table1.cell(23+i,1).text = PtList['PtData'][pt]['LIS'][0]['value'][PtList['PtData'][pt]['LIS'][0]['name'].index(lab_item[i])] 
            
        table1.cell(22,0).text = 'AC BS/ Albumin'
        # ---Infection---
        table1.cell(23,0).text = 'Infection'
        table1.cell(24,0).text = 'WBC/Hgb/PLT'
        table1.cell(25,0).text = 'N/L'
        table1.cell(26,0).text = 'CRP/PCT/Lac'
        table1.cell(27,0).text = 'Na/K/Cl'
        table1.cell(28,0).text = 'BUN/Cr'
        table1.cell(29,0).text = 'TCa/FCa/Mg'
        table1.cell(30,0).text = 'AST/ALT'
        table1.cell(31,0).text = 'PT/INR/PTT'
        table1.cell(32,0).text = 'T/D.bil/TP'
        table1.cell(33,0).text = 'CK/CKMB/TnI'
        table1.cell(34,0).text = 'Amy/Lip/NH3'
        table1.cell(35,0).text = 'LDL/TG/TC/UA'
        table1.cell(36,0).text = 'BNP/D-dimer'
        table1.cell(37,0).text = 'IP/Alk-P/r-GT'
        table1.cell(38,0).text = 'Blood ketone'
        table1.cell(39,0).text = 'Urine/stool'
        # ---LIS---
        table1.cell(2,2).merge(table1.cell(40,2))
        table1.cell(2,2).text = 'n'.join([ PtList['PtData'][pt]['LIS'][0]['name'][x]+'-'+PtList['PtData'][pt]['LIS'][0]['value'][x] for x in range(len(PtList['PtData'][pt]['LIS'][0]['name'])) ])
        # ---Culture---
        INF = []
        for i in range(len(PtList['PtData'][pt]['culture']['access'])) :
            inf_temp = PtList['PtData'][pt]['culture']['collect'][i]+PtList['PtData'][pt]['culture']['type'][i]+'-'
            if PtList['PtData'][pt]['culture']['org'][i] != [] : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['org'][i])
            else : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['other'][i])
            INF.append(inf_temp)
        table1.cell(2,3).add_paragraph('\n'.join(INF))        
        # ---Image---
        table1.cell(2,4).add_paragraph('\n'.join(PtList['PtData'][pt]['RIS'][:min(3,len(PtList['PtData'][pt]['RIS']))-1]))
        # ---CXR近期一張---
        for i in range(min(1,len(PtList['PtData'][pt]['CXR']['date']))) :
            table1.cell(2,4).add_paragraph(PtList['PtData'][pt]['CXR']['date'][i])
            img = sessions.get(PtList['PtData'][pt]['CXR']['url'][i]).content
            r = table1.cell(2,4).add_paragraph().add_run()
            r.add_picture(BytesIO(img))
    print(doc)
    return doc
# ---NS格式---
def NSform(PtList, Ptnum) :
    doc = Document('default.docx')
    for pt in range(Ptnum) :
        print('Pt[',pt+1,']->', end="")
        section = doc.sections[-1]
        style = doc.styles['Normal']
        style.font.size = Pt(9)
        # ---版面配置---
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        # ---新增ptlist表格---
        # ---General data---
        table1 = doc.add_table(rows = 1, cols = 5, style = 'Table Grid')
        table1.allow_autofit = False
        table1.cell(0,0).text = PtList['CHARTNO'][pt]+'\n'+PtList['NRCODE'][pt]+'-'+PtList['BEDNO'][pt]+'\n'+PtList['NameGenderAge'][pt]+'\n'+PtList['DepName'][pt]+'\nVS: '+PtList['VSDRNAME'][pt]+'\n'+PtList['INDATETIME'][pt][:7]
        table1.cell(0,1).text = '\n'.join(PtList['PtData'][pt]['imp'])
        # ---TPR chart---
        table1.cell(0,4).text = 'TPR: '+PtList['PtData'][pt]['MIchart']['T']+'/'+PtList['PtData'][pt]['MIchart']['P']+'/'+PtList['PtData'][pt]['MIchart']['R']
        table1.cell(0,4).add_paragraph('IO: '+PtList['PtData'][pt]['MIchart']['IO']['IT']+'/'+PtList['PtData'][pt]['MIchart']['IO']['OT']+'/'+PtList['PtData'][pt]['MIchart']['IO']['IO'])
        # ---Lab data---
        l = min(2,len(PtList['PtData'][pt]['LIS']))
        data = []
        for i in range(l) :
            table1.cell(0,2).add_paragraph(PtList['PtData'][pt]['LIS'][i]['date'])
            data = [ PtList['PtData'][pt]['LIS'][i]['name'][x]+'-'+PtList['PtData'][pt]['LIS'][i]['value'][x] for x in range(len(PtList['PtData'][pt]['LIS'][i]['name'])) ]
            table1.cell(0,2).add_paragraph('\n'.join(data))
        # ---Anti---
        anti = [ PtList['PtData'][pt]['anti']['name'][x]+'-'+PtList['PtData'][pt]['anti']['dose'][x]+'-'+PtList['PtData'][pt]['anti']['date'][x] for x in range(len(PtList['PtData'][pt]['anti']['name'])) ]
        table1.cell(0,3).text = '\n'.join(anti)
        # ---Culture---
        INF = []
        for i in range(len(PtList['PtData'][pt]['culture']['access'])) :
            inf_temp = PtList['PtData'][pt]['culture']['collect'][i]+PtList['PtData'][pt]['culture']['type'][i]+'-'
            if PtList['PtData'][pt]['culture']['org'][i] != [] : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['org'][i])
            else : inf_temp = inf_temp+','.join(PtList['PtData'][pt]['culture']['other'][i])
            INF.append(inf_temp)
        table1.cell(0,3).add_paragraph('\n'.join(INF))
        # ---Image---
        table1.cell(0,4).add_paragraph('\n'.join(PtList['PtData'][pt]['RIS'][:min(3,len(PtList['PtData'][pt]['RIS']))-1]))
        # ---CXR近期一張---
        for i in range(min(1,len(PtList['PtData'][pt]['CXR']['date']))) :
            table1.cell(0,4).add_paragraph(PtList['PtData'][pt]['CXR']['date'][i])
            img = sessions.get(PtList['PtData'][pt]['CXR']['url'][i]).content
            r = table1.cell(0,4).add_paragraph().add_run()
            r.add_picture(BytesIO(img))
    print(doc)
    return doc
# -----------------------------
# ---函數區---
again = True
ID = input("輸入progress note帳號(含DOC): ").upper()
PW = input("輸入密碼: ")
while again : 
    # -----輸入帳號密碼--------
    Doc = Section = Ward = WardDel = BlockW = BlockBs = BlockBe = ['']
    g = input('選擇表單種類 [G]一般/[M]MI/[2]MI2/[C]CCU/[N]NS: ').upper()
    if g == 'G' :
        print('以下請輸入篩選條件 \n*不需要則按Enter跳過\n*多個選項請用","分隔\n')
        Doc = input("輸入主治帳號(含DOC): ").upper()
        Section = input("輸入科別: ").upper().split(',')
        Ward = input("輸入病房代碼: ").upper().split(',')
        WardDel = input("輸入要剔除的病房代碼: ").upper().split(',')
        s = input('是否指定床段[Y/N]: ').upper()
        if s == 'Y' :
            n = input('床段數量? ')
            for i in range(int(n)) :
                BlockW.append(input('病房:').upper())
                BlockBs.append(input('床號起點: ').upper())
                BlockBe.append(input('床號終點: ').upper())
        form = input('請選擇格式[D]etail / [S]imple: ').upper()
        lab = 'G'
    elif g == 'M' :
        Ward = ['MICU']
        BlockW = ['MICU']; BlockBs = ['MI01']; BlockBe = ['MI15']
        lab = 'I'
        form = 'M'
    elif g == '2' :
        Ward = ['MI2']
        lab = 'I'
        form = 'M'
    elif g == 'C' :
        Ward = ['MICU']
        BlockW = ['MICU']; BlockBs = ['CI01']; BlockBe = ['CI15']
        lab = 'I'
        form = 'M'
    elif g == 'N' :
        Section = ['NS']
        Ward = ['SICU', 'MICU', 'CVSI', 'GICU', 'CVSI', 'BC', 'MI2', 'SNCU']
        lab = 'G'
        form = 'N'

    start = time.time()
    # ------登入progressnote--------
    print('LINK START!!')
    sessions = requests.session()
    requ = sessions.get('http://mobilereport.ndmctsgh.edu.tw/eForm/Account/Login')
    tree = html.fromstring(requ.text)
    token = list(set(tree.xpath('//input[@name="__RequestVerificationToken"]/@value')))
    load = { 
        'login_id' : ID , 
        'password' : PW , 
        '__RequestVerificationToken' : token[0] 
    }
    login_prog = sessions.post('http://mobilereport.ndmctsgh.edu.tw/eForm/Account/Login', data = load)

    PtList = GetPtList(Section, Ward, Doc, WardDel, BlockW, BlockBs, BlockBe, sessions)
    Ptnum = len(PtList['CHARTNO'])
    print('病人總數:'+str(Ptnum))
    # ------報告擷取------
    for pt in range(len(PtList['CHARTNO'])) :
        print('Pt['+str(pt+1)+'] '+PtList['NRCODE'][pt]+'-'+PtList['BEDNO'][pt]+' '+PtList['CHARTNO'][pt]+': ', end="")
        # ---TPR---
        data_TPR = GetTPR(PtList['CHARTNO'][pt], PtList['MEDNO'][pt], PtList['VISITSEQ'][pt], sessions)
        # -----Anti list-------
        anti_list = GetAntiList(PtList['CHARTNO'][pt], PtList['MEDNO'][pt], PtList['VISITSEQ'][pt], sessions)
        # ------BP SpO2 I/O list--------
        IO_list = GetIOList(PtList['CHARTNO'][pt], PtList['MEDNO'][pt], PtList['VISITSEQ'][pt], sessions)
        # -----Lab Data------
        [lab_total, culture] = GetLISList(ID, PtList['CHARTNO'][pt], sessions, lab)
        # ------RIS報告--------
        [RIS_list, CXR_url] = GetRISList(ID, PtList['CHARTNO'][pt], sessions)
        # -------院內診斷---------
        imp = GetImpList(ID, PtList['CHARTNO'][pt], PtList['INDATETIME'][pt][0:7], PtList['HCASENO'][pt], sessions) 
        # ------住院用藥---------
        med = GetMedList(ID, PtList['CHARTNO'][pt], sessions)
        # ------ICU flowsheet------
        MI_chart = GetICUsheet(PtList['CHARTNO'][pt], sessions)
        # ---Data Merge---
        PtList['PtData'][pt] = { 'TPR' : data_TPR, 'anti' : anti_list, 'IO' : IO_list, 'LIS' : lab_total, 'culture' : culture, 'RIS' : RIS_list, 'CXR' : CXR_url, 'imp' : imp, 'med' : med, 'MIchart' : MI_chart }
        print('completed!')
    end = time.time() - start
    print(end)

    # --------doc生成-----------
    print('System call Document generate!')
    docname = 'PtList_'+time.strftime('%Y%m%d_%H%M%S')+'_'+g+'.docx'
    print(docname)

    start = time.time()
    if form == 'D' : detailform(PtList, Ptnum).save(docname)
    elif form == 'S' : simpleform(PtList, Ptnum).save(docname)
    elif form == 'M' : MIform(PtList, Ptnum).save(docname)
    elif form == 'N' : NSform(PtList, Ptnum).save(docname)
    end = time.time() - start
    print(end)
    if input('再來一張[Y/N]? ').upper() == 'N' : again = False
input()

